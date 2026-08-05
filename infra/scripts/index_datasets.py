#!/usr/bin/env python3
"""DEPRECATED for production deploys — manual / debug use only.

The official deploy flow now uses scripts/setup_search_pipeline.py which:
  - Creates canonical schemas (id filterable, vectorizer, source_blob, indexed_at)
  - Sets up Blob → AI Search continuous indexers (Step 6 — managed pipeline
    that replaces this script's one-shot upload behaviour)
  - Sets up Cosmos → AI Search continuous indexer for chat history (Step 5)

The continuous blob indexer handles PDF/DOCX/CSV/JSON parsing natively and
generates content_vector via Azure OpenAI Embedding Skill at index time.
Once configured (via setup_search_pipeline.py Step 6), any blob uploaded to
the container is auto-indexed in <15 min without re-running this script.

This script remains useful for:
  - Force-reingest of a specific container during testing
  - Local development before infra is provisioned
  - Debugging schema/extraction issues on individual blobs

NOT to be invoked from selecting_team_config_and_data.{sh,ps1} anymore —
those scripts now call setup_search_pipeline.py centrally.

Schema creado (idéntico a build_doc_index_schema en setup_search_pipeline.py):
    id (key, filterable, retrievable)
    title (searchable, filterable, retrievable)
    content (searchable, retrievable)
    source_blob (filterable, retrievable)
    indexed_at (filterable, sortable, retrievable)
    content_vector (Collection<Single>, dims=1536, vectorSearchProfile=vector-profile)
    + vectorSearch HNSW + azureOpenAIVectorizer
    + semantic config 'default' (titleField=title, contentFields=[content])

Cada doc se sube con embedding ya generado vía Azure OpenAI.

Usage:
    python index_datasets.py <storage_account> <blob_container> <ai_search_endpoint> [<index_name>]

Env vars requeridas:
    AZURE_OPENAI_ENDPOINT
    AZURE_OPENAI_EMBEDDING_DEPLOYMENT (default: text-embedding-3-small)
"""

import io
import os
import sys
from datetime import datetime, timezone
from typing import List, Optional

import requests
from azure.identity import AzureCliCredential
from azure.search.documents import SearchClient
from azure.storage.blob import BlobServiceClient

# ── Config ──────────────────────────────────────────────────────────

EMBEDDING_DIMS = 1536
SEARCH_API = "2024-07-01"
OPENAI_API = "2024-10-21"

OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "").rstrip("/")
EMBEDDING_DEPLOYMENT = os.getenv(
    "AZURE_OPENAI_EMBEDDING_DEPLOYMENT", "text-embedding-3-small"
)

# ── Schema templates (deben coincidir con scripts/setup_search_pipeline.py) ──

VECTOR_SEARCH_CONFIG = {
    "algorithms": [
        {
            "name": "hnsw-algo",
            "kind": "hnsw",
            "hnswParameters": {
                "m": 4,
                "efConstruction": 400,
                "efSearch": 500,
                "metric": "cosine",
            },
        }
    ],
    "profiles": [
        {
            "name": "vector-profile",
            "algorithm": "hnsw-algo",
            "vectorizer": "aoai-vectorizer",
        }
    ],
    "vectorizers": [
        {
            "name": "aoai-vectorizer",
            "kind": "azureOpenAI",
            "azureOpenAIParameters": {
                "resourceUri": OPENAI_ENDPOINT,
                "deploymentId": EMBEDDING_DEPLOYMENT,
                "modelName": EMBEDDING_DEPLOYMENT,
            },
        }
    ],
}

SEMANTIC_CONFIG = {
    "configurations": [
        {
            "name": "default",
            "prioritizedFields": {
                "titleField": {"fieldName": "title"},
                "prioritizedContentFields": [{"fieldName": "content"}],
                "prioritizedKeywordsFields": [],
            },
        }
    ],
}


def build_index_schema(name: str) -> dict:
    """Schema canónico — debe coincidir con build_doc_index_schema en setup_search_pipeline.py."""
    return {
        "name": name,
        "fields": [
            {
                "name": "id",
                "type": "Edm.String",
                "key": True,
                "searchable": False,
                "filterable": True,
                "retrievable": True,
            },
            {
                "name": "title",
                "type": "Edm.String",
                "searchable": True,
                "filterable": True,
                "retrievable": True,
            },
            {
                "name": "content",
                "type": "Edm.String",
                "searchable": True,
                "filterable": False,
                "retrievable": True,
            },
            {
                "name": "source_blob",
                "type": "Edm.String",
                "searchable": False,
                "filterable": True,
                "retrievable": True,
            },
            {
                "name": "indexed_at",
                "type": "Edm.DateTimeOffset",
                "searchable": False,
                "filterable": True,
                "sortable": True,
                "retrievable": True,
            },
            {
                "name": "content_vector",
                "type": "Collection(Edm.Single)",
                "searchable": True,
                "retrievable": False,
                "stored": True,
                "dimensions": EMBEDDING_DIMS,
                "vectorSearchProfile": "vector-profile",
            },
        ],
        "vectorSearch": VECTOR_SEARCH_CONFIG,
        "semantic": SEMANTIC_CONFIG,
    }


# ── Text extractors (sin cambios) ───────────────────────────────────


def extract_pdf_text(pdf_bytes):
    """Extract text content from PDF bytes using PyPDF2."""
    try:
        import PyPDF2

        pdf_file = io.BytesIO(pdf_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        if pdf_reader.is_encrypted:
            return "PDF_PROTECTED: This PDF document is password-protected or encrypted and cannot be processed."

        text_content = []
        for page in pdf_reader.pages:
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_content.append(page_text)
            except Exception:
                continue

        full_text = "\n".join(text_content).strip()

        protection_indicators = [
            "protected by Microsoft Office",
            "You'll need a different reader",
            "Download a compatible PDF reader",
            "This PDF Document has been protected",
        ]
        if any(indicator.lower() in full_text.lower() for indicator in protection_indicators):
            return "PDF_PROTECTED: This PDF document appears to be protected or encrypted."

        return full_text if full_text else "PDF_NO_TEXT: No readable text content found in PDF."

    except ImportError:
        return "PDF_ERROR: PyPDF2 library not available. Install with: pip install PyPDF2"
    except Exception as e:
        return f"PDF_ERROR: Error reading PDF content: {str(e)}"


def extract_docx_text(docx_bytes):
    """Extract text content from DOCX bytes using python-docx."""
    try:
        from docx import Document

        docx_file = io.BytesIO(docx_bytes)
        doc = Document(docx_file)

        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)

        full_text = "\n".join(text_content).strip()
        return full_text if full_text else "DOCX_NO_TEXT: No readable text content found in DOCX."

    except ImportError:
        return "DOCX_ERROR: python-docx library not available. Install with: pip install python-docx"
    except Exception as e:
        return f"DOCX_ERROR: Error reading DOCX content: {str(e)}"


# ── REST helpers (PUT directo para crear/actualizar índice) ─────────


def _search_token(credential) -> str:
    return credential.get_token("https://search.azure.com/.default").token


def _openai_token(credential) -> str:
    return credential.get_token("https://cognitiveservices.azure.com/.default").token


def put_index_canonical(endpoint: str, name: str, credential) -> None:
    """PUT canonical schema. Aborta con mensaje claro si AI Search rechaza por
    CannotChangeExistingField (índice legacy con campos no compatibles)."""
    url = f"{endpoint}/indexes/{name}?api-version={SEARCH_API}"
    headers = {
        "Authorization": f"Bearer {_search_token(credential)}",
        "Content-Type": "application/json",
    }
    schema = build_index_schema(name)
    resp = requests.put(url, headers=headers, json=schema, timeout=60)
    if resp.status_code in (200, 201, 204):
        print(f"Index '{name}' created/updated with canonical schema.")
        return

    body = resp.text or ""
    if "CannotChangeExistingField" in body or resp.status_code == 400:
        print(
            f"ERROR: AI Search rejected canonical schema for '{name}'.\n"
            f"  HTTP {resp.status_code}: {body[:400]}\n\n"
            f"  Causa probable: el índice ya existe con un schema legacy incompatible.\n"
            f"  Solución: ejecuta scripts/setup_search_pipeline.py primero — "
            f"Step 0 (migrate_doc_indices_to_canonical) hace la migración\n"
            f"  segura (dump → DELETE → recreate canonical → re-upload)."
        )
        sys.exit(1)

    print(f"ERROR: PUT index '{name}' failed: HTTP {resp.status_code} — {body[:400]}")
    sys.exit(1)


def generate_embeddings_batch(texts: List[str], credential) -> Optional[List[List[float]]]:
    """Genera embeddings en lotes de 16 vía Azure OpenAI. None on failure."""
    if not texts:
        return []
    if not OPENAI_ENDPOINT:
        print("ERROR: AZURE_OPENAI_ENDPOINT no está configurado en env.")
        return None

    url = (
        f"{OPENAI_ENDPOINT}/openai/deployments/{EMBEDDING_DEPLOYMENT}"
        f"/embeddings?api-version={OPENAI_API}"
    )
    all_embeddings: List[List[float]] = []
    for i in range(0, len(texts), 16):
        batch = [t[:8000] for t in texts[i : i + 16]]
        headers = {
            "Authorization": f"Bearer {_openai_token(credential)}",
            "Content-Type": "application/json",
        }
        resp = requests.post(url, headers=headers, json={"input": batch}, timeout=60)
        if resp.status_code != 200:
            print(
                f"ERROR: embedding batch {i} failed: HTTP {resp.status_code} — "
                f"{resp.text[:300]}"
            )
            return None
        data = resp.json()
        all_embeddings.extend([d["embedding"] for d in data["data"]])
    return all_embeddings


# ── Main ────────────────────────────────────────────────────────────


def main() -> None:
    if len(sys.argv) < 4:
        print(
            "Usage: python index_datasets.py "
            "<storage_account_name> <blob_container_name> <ai_search_endpoint> "
            "[<ai_search_index_name>]"
        )
        sys.exit(1)

    storage_account_name = sys.argv[1]
    blob_container_name = sys.argv[2]
    ai_search_endpoint = sys.argv[3]
    ai_search_index_name = sys.argv[4] if len(sys.argv) > 4 else "sample-dataset-index"

    if "search.windows.net" not in ai_search_endpoint:
        ai_search_endpoint = f"https://{ai_search_endpoint}.search.windows.net"

    if not OPENAI_ENDPOINT:
        print("ERROR: AZURE_OPENAI_ENDPOINT must be set in env.")
        sys.exit(1)

    credential = AzureCliCredential()

    # 1. Listar blobs
    try:
        blob_service_client = BlobServiceClient(
            account_url=f"https://{storage_account_name}.blob.core.windows.net",
            credential=credential,
        )
        container_client = blob_service_client.get_container_client(blob_container_name)
        print("Fetching files in container...")
        blob_list = list(container_client.list_blobs())
    except Exception as e:
        print(f"Error fetching files: {e}")
        sys.exit(1)

    # 2. Crear/actualizar índice con schema canónico (PUT REST directo)
    print("Creating or updating Azure Search index (canonical schema)...")
    put_index_canonical(ai_search_endpoint, ai_search_index_name, credential)

    # 3. Extraer texto de cada blob
    success_count = 0
    fail_count = 0
    data_list = []
    for idx, blob in enumerate(blob_list, start=1):
        title = blob.name
        for ext in (".csv", ".json", ".pdf", ".docx", ".pptx"):
            title = title.replace(ext, "")
        data = container_client.download_blob(blob.name).readall()
        try:
            print(f"Reading data from blob: {blob.name}...")
            if blob.name.lower().endswith(".pdf"):
                text = extract_pdf_text(data)
            elif blob.name.lower().endswith(".docx"):
                text = extract_docx_text(data)
            else:
                text = data.decode("utf-8")

            data_list.append(
                {
                    "id": str(idx),
                    "title": title,
                    "content": text,
                    "source_blob": blob.name,
                    "indexed_at": datetime.now(timezone.utc).isoformat(),
                }
            )
            success_count += 1
        except Exception as e:
            print(f"Error reading file - {blob.name}: {e}")
            fail_count += 1
            continue

    if not data_list:
        print(
            f"No data to upload. Success: {success_count}, Failed: {fail_count}"
        )
        sys.exit(1)

    # 4. Generar embeddings en lote
    print(f"Generating embeddings for {len(data_list)} docs...")
    contents = [d["content"] or d["title"] for d in data_list]
    embeddings = generate_embeddings_batch(contents, credential)
    if embeddings is None:
        print("ERROR: Embedding generation failed. Aborting upload.")
        sys.exit(1)

    for doc, vector in zip(data_list, embeddings):
        doc["content_vector"] = vector

    # 5. Upload via SDK (push API)
    try:
        print("Uploading documents to the index...")
        search_client = SearchClient(
            endpoint=ai_search_endpoint,
            index_name=ai_search_index_name,
            credential=credential,
        )
        result = search_client.upload_documents(documents=data_list)
        successes = sum(1 for r in result if getattr(r, "succeeded", False))
        failures = len(data_list) - successes
        print(
            f"Uploaded documents. Requested: {len(data_list)}, "
            f"Succeeded: {successes}, Failed: {failures}"
        )
    except Exception as e:
        print(f"Error uploading documents: {e}")
        sys.exit(1)

    print(f"Processing complete. Success: {success_count}, Failed: {fail_count}")


if __name__ == "__main__":
    main()
